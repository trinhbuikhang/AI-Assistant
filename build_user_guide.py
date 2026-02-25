"""
Generate AI_Assistant_UserGuide.docx using python-docx.
Run from ai_assistant folder: python build_user_guide.py
"""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

# --- Paths ---
OUTPUT_PATH = Path(__file__).resolve().parent / "AI_Assistant_UserGuide.docx"

# --- Style helpers ---
def set_paragraph_style(p: Paragraph, font_name: str, size_pt: int, bold: bool = False,
                       color_hex: str | None = None, space_before_pt: int | None = None,
                       space_after_pt: int | None = None, italic: bool = False):
    r = p.runs[0] if p.runs else p.add_run("")
    r.font.name = font_name
    r.font.size = Pt(size_pt)
    r.font.bold = bold
    r.font.italic = italic
    if color_hex:
        r.font.color.rgb = parse_hex_color(color_hex)
    if space_before_pt is not None:
        p.paragraph_format.space_before = Pt(space_before_pt)
    if space_after_pt is not None:
        p.paragraph_format.space_after = Pt(space_after_pt)

def parse_hex_color(hex_str: str):
    from docx.shared import RGBColor
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

def add_note_box(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    run = p.add_run(text)
    run.font.italic = True
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    run.font.color.rgb = parse_hex_color("1a1a1a")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "e8f0fe")
    p._p.get_or_add_pPr().append(shd)
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")  # 3pt
    left.set(qn("w:color"), "2f65cb")
    pBdr.append(left)
    p._p.get_or_add_pPr().append(pBdr)
    return p

def add_warning_box(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    run = p.add_run(text)
    run.font.italic = True
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    run.font.color.rgb = parse_hex_color("1a1a1a")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "fff8e1")
    p._p.get_or_add_pPr().append(shd)
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:color"), "ffa726")
    pBdr.append(left)
    p._p.get_or_add_pPr().append(pBdr)
    return p

def style_table(t: Table, header_fill: str = "2f65cb", zebra_fill: str = "f0f4ff", border_color: str = "cccccc"):
    for row_idx, row in enumerate(t.rows):
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            if row_idx == 0:
                shd.set(qn("w:fill"), header_fill)
            elif row_idx % 2 == 1:
                shd.set(qn("w:fill"), zebra_fill)
            tcPr.append(shd)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if row_idx == 0:
                        run.font.color.rgb = parse_hex_color("ffffff")
                    run.font.size = Pt(11)
                    run.font.name = "Times New Roman"

def add_page_number_footer(section, total_pages_placeholder: int = 1):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Page 1 of {total_pages_placeholder}")
    run.font.size = Pt(9)
    run.font.name = "Times New Roman"
    run.font.color.rgb = parse_hex_color("666666")

def add_header_text(section, text: str):
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = "Arial"
    run.font.color.rgb = parse_hex_color("666666")

def main():
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ---------- Cover page ----------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("[LOGO]")
    r.font.size = Pt(12)
    r.font.color.rgb = parse_hex_color("888888")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("AI Assistant — User Guide")
    r.font.name = "Arial"
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("1a3a6b")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(36)
    r = p.add_run("Internal AI assistant, runs completely offline")
    r.font.name = "Arial"
    r.font.size = Pt(14)
    r.font.color.rgb = parse_hex_color("555555")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Version 1.0")
    r.font.name = "Arial"
    r.font.size = Pt(11)
    r.font.color.rgb = parse_hex_color("555555")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"Date: {date.today().strftime('%B %d, %Y')}")
    r.font.name = "Arial"
    r.font.size = Pt(11)
    r.font.color.rgb = parse_hex_color("555555")

    doc.add_page_break()

    # ---------- Table of contents ----------
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Table of Contents")
    r.font.name = "Arial"
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("1a3a6b")

    toc_items = [
        "1. Introduction",
        "2. Launching the Application",
        "3. Main Interface",
        "4. Basic Usage",
        "5. Reading and Processing Documents",
        "6. Managing Conversation History",
        "7. Settings",
        "8. Troubleshooting",
        "9. Security & Privacy",
        "10. Support Contact",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Bullet").paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ---------- Section 1 — Introduction ----------
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("1. Introduction")
    r.font.name = "Arial"
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("1a3a6b")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("1.1 What is AI Assistant?")
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("2f65cb")

    body = (
        "AI Assistant is your company's own AI helper. It runs entirely on your internal network and does not "
        "send any data to the internet. You can use it to summarize documents, answer questions, and help you "
        "draft text and emails — all in a simple, chat-style window."
    )
    doc.add_paragraph(body).paragraph_format.space_after = Pt(6)
    for run in doc.paragraphs[-1].runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    doc.paragraphs[-1].paragraph_format.line_spacing = 1.3

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("1.2 Requirements")
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("2f65cb")

    req1 = "Your computer must be set up by IT with the AI Assistant and Ollama installed."
    doc.add_paragraph(req1).paragraph_format.space_after = Pt(6)
    for run in doc.paragraphs[-1].runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    doc.paragraphs[-1].paragraph_format.line_spacing = 1.3

    req2 = "You need to be connected to the company's internal network."
    doc.add_paragraph(req2).paragraph_format.space_after = Pt(6)
    for run in doc.paragraphs[-1].runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    doc.paragraphs[-1].paragraph_format.line_spacing = 1.3

    req3 = "No internet connection is required — everything runs offline."
    doc.add_paragraph(req3).paragraph_format.space_after = Pt(6)
    for run in doc.paragraphs[-1].runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    doc.paragraphs[-1].paragraph_format.line_spacing = 1.3

    doc.add_page_break()

    # ---------- Section 2 — Launching ----------
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("2. Launching the Application")
    r.font.name = "Arial"
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("1a3a6b")

    steps = [
        ("Step 1:", "Open the application folder on the network drive where the AI Assistant has been installed."),
        ("", "[📸 Screenshot: open Windows Explorer to network drive]"),
        ("Step 2:", "Double-click the run.bat file to start the application."),
        ("", "[📸 Screenshot: run.bat file]"),
        ("Step 3:", "Wait for the application window to open. The first time may take 10–30 seconds."),
        ("", "[📸 Screenshot: main interface]"),
    ]
    for label, text in steps:
        para = doc.add_paragraph()
        if label:
            r = para.add_run(label + " ")
            r.font.name = "Times New Roman"
            r.font.size = Pt(11)
            r.font.bold = True
        r2 = para.add_run(text)
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(11)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.3

    add_note_box(doc, "If you see \"Ollama is not running\", wait about 30 seconds and then click Retry.")

    doc.add_page_break()

    # ---------- Section 3 — Main Interface ----------
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("3. Main Interface")
    r.font.name = "Arial"
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("1a3a6b")

    diagram = (
        "The window is divided into clear areas. You can think of it like this:\n\n"
        "┌─────────────┬──────────────────────────┐\n"
        "│   SIDEBAR   │         TOP BAR          │\n"
        "│  New Chat   │  Title  | Model  |  ⚙️  │\n"
        "│  Chat list  ├──────────────────────────┤\n"
        "│             │                          │\n"
        "│             │       CHAT AREA          │\n"
        "│             │                          │\n"
        "│             ├──────────────────────────┤\n"
        "│ Clear All   │  📎  Input box       ➤  │\n"
        "└─────────────┴──────────────────────────┘"
    )
    para = doc.add_paragraph(diagram)
    para.paragraph_format.space_after = Pt(6)
    for run in para.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(10)
    para.paragraph_format.line_spacing = 1.3

    items = [
        ("Sidebar (left):", "Shows a list of your saved conversations. Click \"＋ New Chat\" to start a new one."),
        ("Top bar:", "Shows the conversation title, lets you select the AI model, and open Settings (⚙️)."),
        ("Chat area (centre):", "Where your messages and the AI's replies appear."),
        ("Input bar (bottom):", "Type your question here and press Enter or click ➤ to send. Use 📎 to attach a file."),
    ]
    for title, desc in items:
        para = doc.add_paragraph()
        r = para.add_run(title + " ")
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
        r.font.bold = True
        r2 = para.add_run(desc)
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(11)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.3

    doc.add_page_break()

    # ---------- Section 4 — Basic Usage ----------
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("4. Basic Usage")
    r.font.name = "Arial"
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("1a3a6b")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("4.1 Starting a new conversation")
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("2f65cb")

    doc.add_paragraph(
        "Click the \"＋ New Chat\" button. Type your question in the input box at the bottom. Press Enter or click ➤ to send. "
        "The AI's reply will appear gradually — this is normal. [📸 Screenshot: typing a message]",
        style="Normal"
    ).paragraph_format.space_after = Pt(6)
    for run in doc.paragraphs[-1].runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    doc.paragraphs[-1].paragraph_format.line_spacing = 1.3

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("4.2 How to ask effective questions")
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("2f65cb")

    doc.add_paragraph("Examples of good questions:", style="Normal").paragraph_format.space_after = Pt(4)
    for run in doc.paragraphs[-1].runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    good = [
        '"Summarize the main points of this document"',
        '"Explain clause 3 in this contract"',
        '"Write a reply email to a customer about a late delivery"',
    ]
    for g in good:
        doc.add_paragraph("✓ " + g, style="List Bullet").paragraph_format.space_after = Pt(2)
        for run in doc.paragraphs[-1].runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

    doc.add_paragraph("Questions to avoid:", style="Normal").paragraph_format.space_before = Pt(8)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(4)
    for run in doc.paragraphs[-1].runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    bad = [
        "Too short or missing context (e.g. just \"summarize\" without attaching or describing the document)",
        "Real-time information (e.g. \"what is the stock price today\") — the AI does not have live internet data.",
    ]
    for b in bad:
        doc.add_paragraph("✗ " + b, style="List Bullet").paragraph_format.space_after = Pt(2)
        for run in doc.paragraphs[-1].runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("4.3 Continuing a conversation")
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("2f65cb")

    doc.add_paragraph(
        "The AI remembers everything within the same conversation. You can follow up based on previous answers. "
        "For example: \"Can you explain point 2 you just mentioned in more detail?\"",
        style="Normal"
    ).paragraph_format.space_after = Pt(6)
    for run in doc.paragraphs[-1].runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    doc.paragraphs[-1].paragraph_format.line_spacing = 1.3

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("4.4 Stopping the AI mid-response")
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("2f65cb")

    doc.add_paragraph(
        "Click the red \"⏹ Stop generating\" button. The AI will stop immediately.",
        style="Normal"
    ).paragraph_format.space_after = Pt(6)
    for run in doc.paragraphs[-1].runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    doc.add_page_break()

    # ---------- Section 5 — Documents ----------
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("5. Reading and Processing Documents")
    r.font.name = "Arial"
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("1a3a6b")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("5.1 Supported file formats")
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("2f65cb")

    t = doc.add_table(rows=7, cols=3)
    t.style = "Table Grid"
    rows_data = [
        ("Format", "Description", "Supported"),
        (".pdf", "Adobe PDF", "Yes"),
        (".docx", "Microsoft Word", "Yes"),
        (".txt", "Plain text", "Yes"),
        (".csv", "Data table", "Yes"),
        (".xlsx", "Excel spreadsheet", "No"),
        (".jpg / .png", "Images", "No"),
    ]
    for i, row_data in enumerate(rows_data):
        row = t.rows[i]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = cell_text
    style_table(t)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("5.2 How to attach a file")
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("2f65cb")

    attach_steps = [
        "Click the 📎 icon at the bottom right of the input area.",
        "Select a file from your computer.",
        "You will see the filename appear above the input box — that means the file is attached.",
        "Type your question about the file and press Enter. [📸 Screenshot: attached file badge]",
    ]
    for i, step in enumerate(attach_steps, 1):
        doc.add_paragraph(f"Step {i}: {step}", style="List Number").paragraph_format.space_after = Pt(4)
        for run in doc.paragraphs[-1].runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
        doc.paragraphs[-1].paragraph_format.line_spacing = 1.3

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("5.3 Real-world example")
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("2f65cb")

    doc.add_paragraph(
        "Situation: You need to summarize a long PDF report. Attach the report file, then type: "
        "\"Please summarize the key points of this document in 5 bullet points.\" The AI will read the document and respond.",
        style="Normal"
    ).paragraph_format.space_after = Pt(6)
    for run in doc.paragraphs[-1].runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    doc.paragraphs[-1].paragraph_format.line_spacing = 1.3

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("5.4 Notes on large files")
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("2f65cb")

    doc.add_paragraph(
        "For files over about 50 pages, the AI may split the content into chunks and take 1–2 minutes to process. "
        "Files with many images: only the text portions are read; images are not analysed.",
        style="Normal"
    ).paragraph_format.space_after = Pt(6)
    for run in doc.paragraphs[-1].runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    doc.paragraphs[-1].paragraph_format.line_spacing = 1.3

    doc.add_page_break()

    # ---------- Section 6 — History ----------
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("6. Managing Conversation History")
    r.font.name = "Arial"
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("1a3a6b")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("6.1 Viewing past conversations")
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("2f65cb")
    doc.add_paragraph(
        "The list on the left shows all your saved conversations. Click a title to reopen it. You can continue the conversation from where you left off.",
        style="Normal"
    ).paragraph_format.space_after = Pt(6)
    for run in doc.paragraphs[-1].runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    doc.paragraphs[-1].paragraph_format.line_spacing = 1.3

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("6.2 Starting a new conversation")
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("2f65cb")
    doc.add_paragraph(
        "Click \"＋ New Chat\". The previous conversation is saved automatically.",
        style="Normal"
    ).paragraph_format.space_after = Pt(6)
    for run in doc.paragraphs[-1].runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("6.3 Clearing all history")
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("2f65cb")
    doc.add_paragraph(
        "Click \"Clear All History\" at the bottom of the sidebar. Confirm in the dialog box.",
        style="Normal"
    ).paragraph_format.space_after = Pt(6)
    for run in doc.paragraphs[-1].runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    add_warning_box(doc, "Warning: Clearing all history cannot be undone.")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("6.4 Where is history stored?")
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("2f65cb")
    doc.add_paragraph(
        "On your personal computer: C:\\Users\\[your username]\\Documents\\AI_Assistant\\chat_history\\",
        style="Normal"
    ).paragraph_format.space_after = Pt(4)
    for run in doc.paragraphs[-1].runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    doc.add_paragraph(
        "No one else can see your conversation history.",
        style="Normal"
    ).paragraph_format.space_after = Pt(6)
    for run in doc.paragraphs[-1].runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    doc.add_page_break()

    # ---------- Section 7 — Settings ----------
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("7. Settings")
    r.font.name = "Arial"
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("1a3a6b")

    doc.add_paragraph(
        "Open settings by clicking the ⚙️ icon in the top right.",
        style="Normal"
    ).paragraph_format.space_after = Pt(12)
    for run in doc.paragraphs[-1].runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("7.1 Available options")
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("2f65cb")

    opts = [
        ("System Prompt:", "Background instruction given to the AI. Example: \"Always respond in English and keep answers concise.\" Leave blank to use the default."),
        ("Temperature (0.0 – 1.0):", "Near 0.0 = precise, consistent answers. Near 1.0 = more varied, creative answers. Recommended: 0.7 for general office work."),
        ("Max Tokens:", "Limits the length of the AI's response. Default 2048 is sufficient for most use cases."),
    ]
    for name, desc in opts:
        para = doc.add_paragraph()
        r = para.add_run(name + " ")
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
        r.font.bold = True
        r2 = para.add_run(desc)
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(11)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.3

    doc.add_page_break()

    # ---------- Section 8 — Troubleshooting ----------
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("8. Troubleshooting")
    r.font.name = "Arial"
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("1a3a6b")

    t = doc.add_table(rows=7, cols=3)
    t.style = "Table Grid"
    trouble_data = [
        ("Problem", "Cause", "Solution"),
        ("\"Ollama is not running\" message", "Ollama has not started yet", "Wait 30 seconds, click Retry. If still failing, contact IT."),
        ("AI responds very slowly", "Model is loading for the first time", "Wait 1–2 minutes; it will be faster next time."),
        ("Cannot attach a file", "File is open in another program", "Close the file and try again."),
        ("App won't open", "Python not installed or not set up", "Contact IT to set up the environment."),
        ("Response is cut off mid-sentence", "Max tokens too low", "Go to Settings, increase Max Tokens to 4096."),
        ("Text appears very slowly", "Machine is under heavy load", "Close other programs and try again."),
    ]
    for i, row_data in enumerate(trouble_data):
        row = t.rows[i]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = cell_text
    style_table(t)

    doc.add_page_break()

    # ---------- Section 9 — Security ----------
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("9. Security & Privacy")
    r.font.name = "Arial"
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("1a3a6b")

    sec_items = [
        "All data is processed on internal machines only.",
        "No information is sent to the internet.",
        "Chat history is stored only on your personal computer.",
    ]
    for item in sec_items:
        doc.add_paragraph("✓ " + item, style="List Bullet").paragraph_format.space_after = Pt(4)
        for run in doc.paragraphs[-1].runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

    add_warning_box(doc, "As a general best practice, avoid entering passwords or highly sensitive data into any AI tool.")

    # ---------- Section 10 — Support ----------
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("10. Support Contact")
    r.font.name = "Arial"
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = parse_hex_color("1a3a6b")

    doc.add_paragraph(
        "Placeholder — fill in your company IT support details:",
        style="Normal"
    ).paragraph_format.space_after = Pt(12)
    for run in doc.paragraphs[-1].runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    support_table = doc.add_table(rows=4, cols=2)
    support_data = [
        ("Name", "_______________"),
        ("Email", "_______________"),
        ("Extension", "_______________"),
        ("Support hours", "_______________"),
    ]
    for i, (label, value) in enumerate(support_data):
        support_table.rows[i].cells[0].text = label + ":"
        support_table.rows[i].cells[1].text = value
    for row in support_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)

    # ---------- Header and footer ----------
    add_header_text(section, "AI Assistant — User Guide")
    add_page_number_footer(section, 15)  # placeholder total

    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
